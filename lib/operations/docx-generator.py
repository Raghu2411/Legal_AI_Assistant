import json
import sys
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_docx(content_json, output_path):
    """
    Generates a legal playbook DOCX from structured JSON content.
    """
    doc = Document()
    
    # Title
    title = doc.add_heading('Firm Legal Playbook', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Version Info
    p = doc.add_paragraph()
    p.add_run(f'Version: {content_json.get("version", "1.0")}\n').bold = True
    p.add_run(f'Status: {content_json.get("status", "Draft")}\n')
    p.add_run(f'Last Updated: {content_json.get("last_updated", "N/A")}\n')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Content Sections
    sections = content_json.get('sections', [])
    for section in sections:
        doc.add_heading(section.get('title', 'Untitled Section'), level=1)
        
        clauses = section.get('clauses', [])
        for clause in clauses:
            doc.add_heading(clause.get('name', 'Untitled Clause'), level=2)
            doc.add_paragraph(clause.get('text', 'No content available.'))
            
            if clause.get('guidance'):
                p = doc.add_paragraph()
                p.add_run('Guidance: ').bold = True
                p.add_run(clause.get('guidance'))
                p.style = 'Quote'

    # Save
    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python docx-generator.py <json_input_path> <docx_output_path>")
        sys.exit(1)
        
    input_path = sys.argv[1]
    output_path = sys.argv[2]
    
    try:
        with open(input_path, 'r') as f:
            content = json.load(f)
            
        generate_docx(content, output_path)
        print(f"Successfully generated DOCX at {output_path}")
    except Exception as e:
        print(f"Error generating DOCX: {str(e)}")
        sys.exit(1)
